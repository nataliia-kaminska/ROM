import docx

def get_docx_info(file_path):
    doc = docx.Document(file_path)
    
    print("--- HEADINGS & STRUCTURE ---")
    for para in doc.paragraphs:
        if para.style.name.startswith('Heading'):
            print(f"{para.style.name}: {para.text}")
    
    print("\n--- FORMATTING SAMPLES ---")
    if doc.paragraphs:
        p = doc.paragraphs[0]
        print(f"Font Name: {p.style.font.name}")
        print(f"Font Size: {p.style.font.size}")
        
    print("\n--- MARGINS ---")
    section = doc.sections[0]
    print(f"Top Margin: {section.top_margin.mm} mm")
    print(f"Bottom Margin: {section.bottom_margin.mm} mm")
    print(f"Left Margin: {section.left_margin.mm} mm")
    print(f"Right Margin: {section.right_margin.mm} mm")

    print("\n--- FULL TEXT (First 2000 chars) ---")
    full_text = "\n".join([p.text for p in doc.paragraphs])
    print(full_text[:2000])

if __name__ == "__main__":
    get_docx_info("Дипломна Максим.docx")
